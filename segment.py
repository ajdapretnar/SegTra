import fnmatch
import os
import re

import docx
import pandas as pd
pd.options.mode.chained_assignment = None  # default='warn'

from similarity import aq_distance, chunk_distance, aq_cosine, chunk_cosine

class Reader:
    def __init__(self, document):
        self.document = document

    def question(self, paragraph):
        raise NotImplementedError

    def join_q_and_a(self):
        questions = []
        answers = []
        a = []
        for i, row in enumerate(self.document.paragraphs):
            if self.question(row):
                if a:
                    answers.append("\n".join(a))
                a = []
                questions.append(row.text)
            else:
                a.append(row.text)
        if a:
            answers.append("\n".join(a))
        return questions, answers


class NameReader(Reader):
    def __init__(self, document):
        super().__init__(document)
        self.q_label = self.determine_label(self.document)
        self.document = self.remove_empty_paragraphs(document)

    def question(self, paragraph):
        return paragraph.text.startswith(self.q_label)

    def determine_label(self, document):
        regex = '\w+:'
        return re.match(regex, document.paragraphs[0].text).group()

    def remove_empty_paragraphs(self, document):
        new_doc = docx.Document()
        for paragraph in document.paragraphs:
            if paragraph.text:
                # add just text (does not retain styling)
                new_doc.add_paragraph(paragraph.text)
        return new_doc

class ItalicReader(Reader):
    def __init__(self, document):
        super().__init__(document)

    def question(self, paragraph):
        for run in paragraph.runs:
            if run.italic:
                return True
        return

class BoldReader(Reader):
    def __init__(self, document):
        super().__init__(document)

    def question(self, paragraph):
        for run in paragraph.runs:
            if run.bold:
                return True
        return

class ListReader(Reader):
    def __init__(self, document):
        super().__init__(document)

    def question(self, paragraph):
        # returns True if the paragraph is a list item
        return paragraph._p.pPr.numPr is not None


READERS = {'NameReader': NameReader, 'ItalicReader': ItalicReader,
           'BoldReader': BoldReader, 'ListReader': ListReader}

class Segmenter:
    def __init__(self):
        self.corpus = None
        self.remove = None
        self.seg_corpus = None

    def get_paths(self, folder):
        paths = []
        for dirpath, dirnames, filenames in os.walk(folder):
            for file in list(filenames):
                if file.endswith(".docx") and not (fnmatch.fnmatch(file, ".*")
                                                   or file.startswith("~$")):
                    paths.append(os.path.join(dirpath, file))
                else:
                    continue
        return paths

    def sniff_type(self, document):
        "Check first five paragraphs in a document to determine the reader."
        regex = '\w+:'
        reader = {'NameReader': 0, 'ItalicReader': 0, 'BoldReader': 0,
                  'ListReader': 0}
        for i in range(min(5, len(document.paragraphs))):
            if re.match(regex, document.paragraphs[i].text):
                reader['NameReader'] += 1
            i = False
            b = False
            for run in document.paragraphs[i].runs:
                if run.italic:
                    i = True
                if run.bold:
                    b = True
            if i:
                reader['ItalicReader'] += 1
            if b:
                reader['BoldReader'] += 1
            if document.paragraphs[i].style.name == 'List Paragraph':
                reader['ListReader'] += 1
        if max(reader) == 0:
            raise TypeError("File cannot be read.")
        else:
            return READERS[max(reader.keys(), key=(lambda k: reader[k]))]

    def read(self, folder, remove=False):
        self.remove = remove
        paths = self.get_paths(folder)
        corpus = {'File': [], 'Questions': [], 'Answers': []}
        for path in paths:
            with open(path, 'rb') as f:
                doc = docx.Document(f)
                reader = self.sniff_type(doc)
                try:
                    print(path)
                    r = reader(doc)
                except:
                    n = os.path.basename(path)
                    TypeError(f"{n} cannot be read.")
                    continue
                q, a = r.join_q_and_a()
                for q, a in zip(q, a):
                    corpus['File'].append(path)
                    corpus['Questions'].append(q)
                    corpus['Answers'].append(a)
        self.corpus = pd.DataFrame.from_dict(corpus)

    def segment(self, method):
        files = sorted(set(self.corpus['File']))
        seg_corpus = []
        for i in files:
            subset = self.corpus[self.corpus['File'] == i]
            seg_corpus.append(method(subset))
        seg_corpus = pd.concat(seg_corpus)
        self.seg_corpus = seg_corpus

    def save_data(self, path):
        self.seg_corpus.to_csv(path, index=False)


s = Segmenter()
s.read('path/to/folder', remove=False)
s.segment(aq_distance)
s.save_data('segmented.csv')
